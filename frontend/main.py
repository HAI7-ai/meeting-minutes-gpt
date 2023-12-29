""" A Streamlit page that takes the documents and provide an interface to query with the data available in the document.
"""
import os
import sys
import time
import json
import csv
import pandas as pd
import streamlit as st
import PyPDF2
from docx import Document
from azure.storage.blob import BlobServiceClient
from pages.settings import (
    page_config,
    custom_css,
    delete_folder_contents,
    write_uploaded_files,
    meeting_minutes
)

# Get the absolute path to the project root directory
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
src_path = os.path.abspath(os.path.join(project_root, "src"))
sys.path.insert(0, src_path)

# Load the config.json file
with open(f"{project_root}/config/config.json", "r") as config_file:
    config = json.load(config_file)

# Load Config Values
KNOWLEDGE_BASE_DIR = config[
    "KNOWLEDGE_BASE_DIR"
]  # Load Knowledge base directory name
FAISS_DB_DIR = config["FAISS_DB_DIR"]  # Load Vector database directory name

# Loading prompt templates and GPT Utilities from src
from prompts import prompt_doc_qa
from db_utils import VECTOR_DB_UTILS

# Initialize Vector database
vector_db = VECTOR_DB_UTILS()

# Path for the knowledge base documents
kb_path = f"{project_root}/{KNOWLEDGE_BASE_DIR}"
db_path = f"{project_root}/{FAISS_DB_DIR}"
processed_dir_path = f"{project_root}/processed_documents"
current_db_info_file_path = f"{project_root}/db_details.csv"
processed_files_path = f"{project_root}/processed_files.txt"
uploaded_files_path = f"{project_root}/uploaded_files.txt"
mm_uploads_path = f"{project_root}/uploads"


if "db_exist" not in st.session_state:
    st.session_state.db_exist = False
    st.session_state.db_list = False

connection_string = os.environ.get("CONNECTION_STRING")
container_name = "meeting-minutes"

blob_service_client = BlobServiceClient.from_connection_string(connection_string)
container_client = blob_service_client.get_container_client(container_name)

def process_documents(merge_with_exist: bool=True):
    """ A streamlit function to convert the uploaded document files into chunks and store in vector db.
    """
    try:
        db, db_build_time = vector_db.run_db_build(input_type="documents", embeddings=st.session_state.gpt.embeddings, merge_with_existing_db=merge_with_exist)
        if db is not None:
           # st.info(f"Database build completed in {db_build_time:.4f} seconds")
            st.session_state.db_exist = True
            return st.session_state.db_exist
        else:
            st.session_state.db_exist = False
            return st.session_state.db_exist

    except Exception as e:
        error_msg = f"An error occurred while reading files: {e}"
        st.error(error_msg)
        st.session_state.db_exist = False
        return st.session_state.db_exist  

def input_documents():
    """ A streamlit function to provide upload interface for documents and extract information from it.
    """

    with st.form("Process_Documents", clear_on_submit=True):
        uploaded_files = st.file_uploader(label="Choose files",
                                   type=["pdf", "docx"],
                                   accept_multiple_files=True,
                                   disabled=not st.session_state.valid_key,)
        submit_button = st.form_submit_button(label="Process", disabled=not st.session_state.valid_key)

        if submit_button:
            # Upload all the documents to a temporary directory
            upload_state = write_uploaded_files(uploaded_files=uploaded_files, folder_path=kb_path)
            if not upload_state:
                st.error("Error while uploading files. Please check input files.")
            else:
                with st.spinner("Building database..."):
                    db_status = process_documents(True)
                    st.session_state.db_list = True
                    
    return uploaded_files           

def query_form():
    """A streamlit function to define the query form and subsequent results upon submitting the query."""

    response = None

    with st.form("QnA_Data"):
        input_query = st.text_input(
            label="Please enter the question that can be answered by the uploaded transcripts.",
            placeholder="Enter your question",
        )
        return_source_docs = st.toggle(label="return source documents info", value=False)
        submit_query = st.form_submit_button(
            label="Submit Query", disabled=not st.session_state.valid_key
        )

    if submit_query:
        start_time = time.time()
        local_db = vector_db.load_local_db(embeddings=st.session_state.gpt.embeddings)
        if local_db is not None:
            with st.spinner("Retrieving response ..."):
                response = st.session_state.gpt.retrieval_qa(
                    query=input_query,
                    prompt=prompt_doc_qa(),
                    db=local_db,
                    return_source_documents=return_source_docs,
                )
        else:
            st.error("Database does not exist. Please build the database first.")
            
        end_time = time.time()
        
    if response is not None:
        response_completion = response["result"]
        
        response_source_docs = []
        
        if return_source_docs:
            source_docs = response["source_documents"]
            
            for document in source_docs:
                response_source_docs.append(
                    {
                        "source": document.metadata["source"],
                        "content": document.page_content,
                    }
                )

        with st.expander("", expanded=True):
            st.markdown(response_completion)
            
        if return_source_docs:
            st.markdown(
                f"<p style='font-size: smaller; color: green;'>Source documents: {response_source_docs}</p>",
                unsafe_allow_html=True,
            )
            
        st.markdown(
            f"<p style='font-size: smaller; color: green;'>Reponse time: {(end_time - start_time):.4f} seconds</p>",
            unsafe_allow_html=True,
        )
        
def copy_and_process_files(kb_path, processed_file):
    try:
        if not os.path.exists(kb_path):
            os.makedirs(kb_path)
        
        new_docx_files = []
        
        if not os.path.exists(processed_file):
            open(processed_file, 'w').close()
        
        with open(processed_file, 'r') as file:
            processed_files = file.read().splitlines()
            
        blobs_list = container_client.list_blobs()
        
        for blob in blobs_list:
            if blob.name.lower().endswith(('.docx', '.pdf')):
                if blob.name not in processed_files:
                    new_docx_files.append(blob.name) 
                    
        for blob.name in new_docx_files:
            blob_client = container_client.get_blob_client(blob)
            destination_file_path = os.path.join(kb_path, blob.name)
            with open(destination_file_path, "wb") as local_file:
                download_stream = blob_client.download_blob()
                local_file.write(download_stream.readall())
            print(f"Copied '{blob.name}' to '{kb_path}'")          
            with open(processed_file, 'a') as file:
                file.write(blob.name + '\n')
            process_documents(True)
    
    except Exception as e:
        print(f"An error occurred: {e}")

def delete_files_in_folder(folder_path, uploaded_files_txt):
    try:
        # Read filenames from uploaded_files.txt
        with open(uploaded_files_txt, 'r') as file:
            uploaded_files = file.readlines()
            uploaded_files = [file.strip() for file in uploaded_files]

        for file_name in uploaded_files:
            file_path = os.path.join(folder_path, file_name)
            
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"File '{file_name}' has been deleted.")
            else:
                print(f"File '{file_name}' does not exist in the specified folder.")
                
    except Exception as e:
        print(f"An error occurred: {e}")

def delete_rows_from_csv(csv_file, uploaded_files_txt):
    try:
        # Validate if the CSV file exists
        if not os.path.isfile(csv_file):
            print("CSV file not found.")
            return

        # Read the filenames from the uploaded_files.txt
        with open(uploaded_files_txt, 'r') as file:
            uploaded_files = file.readlines()
            uploaded_files = [file.strip() for file in uploaded_files]

        # Read the CSV file into a list of dictionaries
        with open(csv_file, 'r', newline='') as file:
            reader = csv.DictReader(file)
            rows = list(reader)

        # Filter rows based on 'File_Name' column
        updated_rows = filter(lambda row: row['File_Name'] not in uploaded_files, rows)
        
        # Define field names for writing back to CSV
        fieldnames = ['Input_Type', 'File_Name', 'File_Type', 'Executed_Time']

        if updated_rows:
            # Write updated content back to the CSV file
            with open(csv_file, 'w', newline='') as file:
                writer = csv.DictWriter(file, fieldnames=fieldnames)
                writer.writeheader()  # Write header
                writer.writerows(updated_rows)  # Write updated rows

                print("Filenames deleted successfully from the CSV file.")

            # Clear the contents of uploaded_files.txt
            with open(uploaded_files_txt, 'w') as file:
                file.truncate(0)

    except FileNotFoundError:
        print("CSV file or uploaded files text not found.")
    except PermissionError:
        print("Permission denied: Unable to write to the CSV file.")
    except Exception as e:
        print(f"An error occurred: {e}")

def read_text(file_path):
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    elif file_path.endswith('.pdf'):
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            text = ""
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]  # Access pages as an attribute
                text += page.extract_text()
            return text
    else:
        return "Unsupported file format"

def save_as_docx(minutes, filename):   
    filename_without_extension = os.path.splitext(filename)[0]
    doc = Document()  
    
    # Get the user's download folder
    download_folder = os.path.expanduser("~") + "/Downloads/"
    
    # Combine the download folder path with the filename and .docx extension
    file_path = os.path.join(download_folder, "Meeting_Minutes_" + filename_without_extension + ".docx")
    
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
        
    doc.save(file_path)
    
    # Check if the file was successfully saved
    if os.path.exists(file_path):
        st.success(f"File saved successfully! You can find it in your Downloads folder as '{file_path}'.")

    else:
        st.error("There was an issue saving the file. Please try again.")
    
    delete_folder_contents(mm_uploads_path)

def upload_files():
    """ A streamlit function to provide upload interface for documents and extract information from it.
    """

    with st.form("Upload Files", clear_on_submit=True):
        uploaded_file = st.file_uploader(
            label="Choose files",
            type=["pdf", "docx"],
            accept_multiple_files=False,
            disabled=not st.session_state.valid_key,
        )
        submit_button = st.form_submit_button(
            label="Generate MoM", disabled=not st.session_state.valid_key
        )

        if submit_button:
            if uploaded_file is None:
                st.warning("Please upload a file.")
            else:
                with st.spinner("Please wait. Generating Minutes of Meeting..."):
                    # Create the directory if it doesn't exist
                    if not os.path.exists(mm_uploads_path):
                        os.makedirs(mm_uploads_path)
                    file_path = os.path.join(mm_uploads_path, uploaded_file.name)
                    with open(file_path, "wb") as f:
                        f.write(uploaded_file.read())
                                           
                    transcription = read_text(file_path)
                    minutes = meeting_minutes(transcription)
                    filename = os.path.basename(file_path)
                    save_as_docx(minutes, filename)

def main_page():
    """Streamlit content for Admin page"""

    st.info(
        """
        **Meeting Minutes** uses LLM technology to quickly analyse transcripts of meetings that have been uploaded. With its ability to extract important insights seamlessly, highlight actionable points, and provide thorough summaries, it gives users the power to navigate meetings effectively and derive value from their conversations.
        """
    )
    
    ask_questions_tab, upload_transcripts_tab, meeting_minutes_tab = st.tabs(["Ask Questions", "Upload Transcripts", "Meeting Minutes"])
    
    with ask_questions_tab: 
        query_form()
        
    with upload_transcripts_tab:
        uploaded_files = input_documents()
        if uploaded_files and len(uploaded_files) > 0:
            existing_files = set()
            # Check existing file names
            if os.path.exists("uploaded_files.txt"):
                with open("uploaded_files.txt", "r") as existing_file:
                    for line in existing_file:
                        existing_files.add(line.strip())

            with open("uploaded_files.txt", "a") as file:  # Use "a" for append mode
                for file_obj in uploaded_files:
                    file_name = file_obj.name if hasattr(file_obj, 'name') else os.path.basename(file_obj)
                    if file_name not in existing_files:
                        file.write(f"{file_name}\n")
                        existing_files.add(file_name)
                    else:
                        print(f"File '{file_name}' already exists in uploaded_files.txt and was not appended.")
 
        refresh_database = st.button(label="Refresh", key="Refresh", use_container_width=False)   
        st.caption('_:blue[* **Refresh** - downloads and processess the transcripts available in azure blob storage]_')       
        if refresh_database:
            copy_and_process_files(kb_path, processed_files_path) 
            st.toast('Database has been refreshed!')
            time.sleep(.5)
            
        # delete_database = st.button(label="Delete Transcripts", key="Delete", use_container_width=False)
        # if delete_database:
        #     delete_folder_contents(kb_path)
        #     delete_folder_contents(db_path)
        #     delete_folder_contents(processed_dir_path)
        #     if st.session_state.db_list:
        #         os.remove(current_db_info_file_path)
        #         with open(processed_files_path, 'w') as file:
        #             file.truncate(0)  # This truncates the file, removing all content
        #     st.session_state.db_list = False
        #     st.toast('Database has been deleted!')  
        # st.caption('_:red[* **Delete Transcripts** - deletes all the transcripts from the database]_')  
        
        drop_database = st.button(label="Delete", use_container_width=False)
        st.caption('_:red[* **Delete** - deletes only the uploaded transcripts]_') 
        # if drop_database:
        #     delete_files_in_folder(processed_dir_path, uploaded_files_path)
        #     delete_rows_from_csv(current_db_info_file_path, uploaded_files_path)
        #     st.toast('Uploaded transcripts has been deleted!')
        #     time.sleep(.5) 
        
        # drop_database = st.button(label="Delete", use_container_width=False)
        # st.caption('_:red[* **Delete** - deletes only the uploaded transcripts]_') 

        if drop_database:
            if os.path.exists(processed_dir_path) and os.path.exists(uploaded_files_path):
                delete_files_in_folder(processed_dir_path, uploaded_files_path)

            if os.path.exists(current_db_info_file_path) and os.path.exists(uploaded_files_path):
                delete_rows_from_csv(current_db_info_file_path, uploaded_files_path)
                st.toast('Uploaded transcripts has been deleted!')
                time.sleep(.5) 

        st.session_state.db_list = os.path.exists(current_db_info_file_path)
        if st.session_state.db_list:
            documents_df = pd.read_csv(current_db_info_file_path)
            st.dataframe(documents_df['File_Name'], hide_index=True, column_config={'File_Name': "Available Transcripts"})

    with meeting_minutes_tab: 
        st.caption('_:blue[Generate meeting minutes including a summary, key points, and action items from the uploaded transcripts]_')
        upload_files()

# Load the page config and custom css from settings
page_config()
custom_css()

main_page()